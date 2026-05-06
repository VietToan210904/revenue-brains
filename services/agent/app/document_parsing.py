import os
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.errors import DocumentProcessingError


@dataclass(frozen=True)
class TextSource:
    text: str
    page_number: int | None = None
    paragraph_index: int | None = None
    line_start: int | None = None
    line_end: int | None = None
    char_start: int | None = None
    char_end: int | None = None


@dataclass(frozen=True)
class ParsedDocument:
    path: Path
    text: str
    sources: list[TextSource]


TEXT_EXTENSIONS = {".txt", ".md", ".markdown"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}


def get_upload_root() -> Path:
    configured_path = os.getenv("UPLOAD_STORAGE_PATH", "./uploads")
    configured = Path(configured_path)

    if configured.is_absolute():
        return configured

    repo_root = Path(__file__).resolve().parents[3]
    repo_candidate = repo_root / configured
    cwd_candidate = Path.cwd() / configured

    if repo_candidate.exists() or not cwd_candidate.exists():
        return repo_candidate

    return cwd_candidate


def resolve_storage_key(storage_key: str) -> Path:
    normalized = storage_key.replace("\\", "/").strip("/")
    parts = [part for part in normalized.split("/") if part]

    if not parts or any(part in {".", ".."} for part in parts):
        raise DocumentProcessingError(
            "invalid_storage_key",
            "The file storage key is invalid.",
            status_code=400,
            details={"fileStorageKey": storage_key},
        )

    return get_upload_root().joinpath(*parts)


def parse_document(storage_key: str, original_filename: str, content_type: str) -> ParsedDocument:
    path = resolve_storage_key(storage_key)

    if not path.exists() or not path.is_file():
        raise DocumentProcessingError(
            "missing_file",
            "The uploaded document file was not found in private storage.",
            status_code=404,
            details={"fileStorageKey": storage_key},
        )

    extension = Path(original_filename or path.name).suffix.lower() or path.suffix.lower()

    if extension in TEXT_EXTENSIONS or content_type in {"text/plain", "text/markdown"}:
        parsed = parse_text_file(path)
    elif extension in PDF_EXTENSIONS or content_type == "application/pdf":
        parsed = parse_pdf_file(path)
    elif (
        extension in DOCX_EXTENSIONS
        or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        parsed = parse_docx_file(path)
    else:
        raise DocumentProcessingError(
            "unsupported_format",
            "This document format is not supported by the current text-based parser.",
            status_code=415,
            details={"contentType": content_type, "filename": original_filename},
        )

    if not parsed.text.strip():
        raise DocumentProcessingError(
            "empty_document_text",
            "No readable text was found in the document.",
            status_code=422,
            details={"contentType": content_type, "filename": original_filename},
        )

    return parsed


def parse_text_file(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    sources = []
    char_position = 0

    for line_number, line in enumerate(text.splitlines(), start=1):
        line_start = char_position
        char_position += len(line) + 1
        if line.strip():
            sources.append(
                TextSource(
                    text=line.strip(),
                    line_start=line_number,
                    line_end=line_number,
                    char_start=line_start,
                    char_end=line_start + len(line),
                )
            )

    return ParsedDocument(path=path, text=text, sources=sources)


def parse_pdf_file(path: Path) -> ParsedDocument:
    try:
        reader = PdfReader(path)
        page_sources = []
        page_texts = []

        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                cleaned = text.strip()
                page_texts.append(cleaned)
                page_sources.append(TextSource(text=cleaned, page_number=index))

        return ParsedDocument(path=path, text="\n\n".join(page_texts), sources=page_sources)
    except Exception as exc:  # noqa: BLE001
        raise DocumentProcessingError(
            "parse_failed",
            "The PDF could not be parsed as text.",
            status_code=422,
            details={"filename": path.name},
        ) from exc


def parse_docx_file(path: Path) -> ParsedDocument:
    try:
        document = DocxDocument(path)
        paragraph_texts = []
        sources = []

        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                paragraph_texts.append(text)
                sources.append(TextSource(text=text, paragraph_index=index))

        return ParsedDocument(path=path, text="\n\n".join(paragraph_texts), sources=sources)
    except Exception as exc:  # noqa: BLE001
        raise DocumentProcessingError(
            "parse_failed",
            "The DOCX file could not be parsed as text.",
            status_code=422,
            details={"filename": path.name},
        ) from exc
